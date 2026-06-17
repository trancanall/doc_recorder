"""
compile_and_word.py
Sinh file .docx từ AI_OUTPUT.

Có thể chạy trực tiếp:
  python compile_and_word.py

Hoặc gọi từ FastAPI:
  from compile_and_word import build_docx
  build_docx(ai_output="...", output_path=Path("output/HUONG_DAN.docx"))
"""

from __future__ import annotations

import os
import re
import struct
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

# ── Paths (dùng khi chạy standalone) ────────────────────────────────────────
_BASE_DIR   = Path(os.getenv("OUTPUT_DIR", Path(__file__).parent / "output"))
_OUTPUT_DIR = _BASE_DIR if _BASE_DIR.is_absolute() else Path(__file__).parent / "output"

# Khi chạy standalone, đặt AI_OUTPUT ở đây
AI_OUTPUT = ""

# ── Colors & fonts ───────────────────────────────────────────────────────────
COLOR_TITLE  = RGBColor(0x1A, 0x52, 0x76)
COLOR_BODY   = RGBColor(0x33, 0x33, 0x33)
COLOR_NOTE   = RGBColor(0x77, 0x77, 0x77)
COLOR_BORDER = "AED6F1"
FONT_NAME    = "Times New Roman"


# ── Regex helpers ────────────────────────────────────────────────────────────

def normalize_image_path(raw: str) -> str:
    if not raw:
        return ""
    raw = raw.strip().strip("`").strip('"\'').replace("\\", "/").lstrip("./")
    if re.match(r"^[A-Za-z]:/", raw):
        p = Path(raw)
        return f"highlighted/{p.name}" if p.parent.name.lower() == "highlighted" else p.name
    if "highlighted/" in raw:
        return f"highlighted/{raw.split('highlighted/')[-1]}"
    if raw.lower().endswith((".png", ".jpg", ".jpeg", ".webp")):
        return f"highlighted/{raw}"
    return raw


def resolve_image_file(md_img_path: str, output_dir: Path) -> Path:
    md_img_path = md_img_path.replace("\\", "/").strip()
    if re.match(r"^[A-Za-z]:/", md_img_path):
        return Path(md_img_path)
    return output_dir / md_img_path


def extract_image_path(text: str) -> str | None:
    patterns = [
        r"(?:Khung\s*hình\s*gốc|Hình\s*ảnh\s*minh\s*họa|Screenshot|Image|Ảnh)\s*\*{0,2}\s*:\s*\*{0,2}\s*`?([^`\n\r]+?\.(?:png|jpg|jpeg|webp))`?",
        r"`([^`\n\r]+?\.(?:png|jpg|jpeg|webp))`",
        r"([A-Za-z]:[\\/][^\n\r]+?\.(?:png|jpg|jpeg|webp))",
        r"((?:\.\/)?(?:output\/)?highlighted\/[^\s`\n\r]+?\.(?:png|jpg|jpeg|webp))",
    ]
    for pat in patterns:
        m = re.search(pat, text, re.IGNORECASE)
        if m:
            return normalize_image_path(m.group(1))
    return None


def extract_pixel(text: str) -> str | None:
    patterns = [
        r"<pixel>\s*\[(\d+)\s*,\s*(\d+)\s*,\s*(\d+)\s*,\s*(\d+)\s*]\s*</pixel>",
        r"Vị\s*trí\s*tương\s*tác\s*\*{0,2}\s*:\s*\*{0,2}\s*`?\[?(\d+)\s*,\s*(\d+)\s*,\s*(\d+)\s*,\s*(\d+)]?`?",
        r"pixel\s*\*{0,2}\s*:\s*\*{0,2}\s*`?\[?(\d+)\s*,\s*(\d+)\s*,\s*(\d+)\s*,\s*(\d+)]?`?",
    ]
    for pat in patterns:
        m = re.search(pat, text, re.IGNORECASE)
        if m:
            return f"[{m.group(1)}, {m.group(2)}, {m.group(3)}, {m.group(4)}]"
    return None


def split_steps(ai_text: str) -> list[tuple[str, str]]:
    pattern = re.compile(
        r"(?im)^\s*(?:#{1,6}\s*)?(?:\*\*)?\s*((?:Bước|Step)\s*\d+\s*[:.\-–]\s*.+?)(?:\*\*)?\s*$"
    )
    matches = list(pattern.finditer(ai_text))
    steps = []
    for idx, m in enumerate(matches):
        title   = m.group(1).replace("**", "").strip()
        start   = m.end()
        end     = matches[idx + 1].start() if idx + 1 < len(matches) else len(ai_text)
        content = ai_text[start:end].strip()
        steps.append((title, content))
    return steps


def clean_meta_lines(text: str) -> str:
    lines = []
    for line in text.splitlines():
        s = line.strip()
        if not s:
            continue
        if re.search(r"Khung\s*hình\s*gốc|Hình\s*ảnh\s*minh\s*họa|Screenshot|Image|Ảnh", s, re.IGNORECASE):
            continue
        if re.search(r"Vị\s*trí\s*tương\s*tác|<pixel>|pixel\s*:", s, re.IGNORECASE):
            continue
        if re.fullmatch(r"-{3,}", s):
            continue
        lines.append(line.rstrip())
    return "\n".join(lines).strip()


# ── Style helpers ────────────────────────────────────────────────────────────

def set_run_font(run, size_pt: float, bold=False, italic=False,
                 color: RGBColor | None = None, mono=False):
    run.font.name = "Courier New" if mono else FONT_NAME
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def add_paragraph_border_bottom(para, color_hex: str = "DDDDDD", size: int = 4):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_paragraph_border_left(para, color_hex: str = "AED6F1", size: int = 12):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(size))
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), color_hex)
    pBdr.append(left)
    pPr.append(pBdr)


def set_paragraph_indent(para, left_cm: float = 0.0):
    pPr = para._p.get_or_add_pPr()
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), str(int(left_cm * 567)))
    pPr.append(ind)


def set_paragraph_spacing(para, before_pt: float = 0, after_pt: float = 6):
    fmt = para.paragraph_format
    fmt.space_before = Pt(before_pt)
    fmt.space_after  = Pt(after_pt)


def add_inline_runs(para, text: str, base_size: float = 11, base_color: RGBColor = COLOR_BODY):
    token_re = re.compile(r"\*\*(.+?)\*\*|`(.+?)`")
    last = 0
    for m in token_re.finditer(text):
        if m.start() > last:
            run = para.add_run(text[last:m.start()])
            set_run_font(run, base_size, color=base_color)
        if m.group(1) is not None:
            run = para.add_run(m.group(1))
            set_run_font(run, base_size, bold=True, color=base_color)
        elif m.group(2) is not None:
            run = para.add_run(m.group(2))
            set_run_font(run, base_size - 1, mono=True, color=COLOR_NOTE)
        last = m.end()
    if last < len(text):
        run = para.add_run(text[last:])
        set_run_font(run, base_size, color=base_color)


def get_image_size(path: Path) -> tuple[int, int]:
    data = path.read_bytes()
    ext  = path.suffix.lower()
    if ext == ".png" and len(data) > 24:
        return struct.unpack(">I", data[16:20])[0], struct.unpack(">I", data[20:24])[0]
    if ext in (".jpg", ".jpeg"):
        i = 2
        while i < len(data) - 8:
            if data[i] == 0xFF:
                marker = data[i + 1]
                if 0xC0 <= marker <= 0xC3:
                    return struct.unpack(">H", data[i + 7:i + 9])[0], struct.unpack(">H", data[i + 5:i + 7])[0]
                i += 2 + struct.unpack(">H", data[i + 2:i + 4])[0]
            else:
                i += 1
    return 800, 450


def add_scaled_image(doc: Document, image_path: Path, max_width_cm: float = 14.0):
    if not image_path.exists():
        print(f"⚠️  Không tìm thấy ảnh: {image_path}")
        return
    w_px, h_px = get_image_size(image_path)
    ratio  = h_px / w_px if w_px else 1
    draw_w = min(Cm(max_width_cm).inches, Inches(w_px / 96).inches)
    draw_h = draw_w * ratio
    para   = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.add_run().add_picture(str(image_path), width=Inches(draw_w), height=Inches(draw_h))
    set_paragraph_spacing(para, after_pt=10)


# ── Public API ───────────────────────────────────────────────────────────────

def build_docx(
    ai_output: str | None = None,
    output_path: Path | None = None,
) -> Path:
    """
    Sinh file .docx.

    Parameters
    ----------
    ai_output   : nội dung AI text (dùng global AI_OUTPUT nếu None)
    output_path : đường dẫn file output (dùng default nếu None)
    """
    text = (ai_output or AI_OUTPUT).strip()
    if not text:
        raise ValueError("ai_output rỗng — không có gì để compile.")

    out = output_path or (_OUTPUT_DIR / "HUONG_DAN_SU_DUNG_FINAL.docx")
    out.parent.mkdir(parents=True, exist_ok=True)

    # Output dir để resolve ảnh — lấy từ out.parent
    img_base = out.parent

    steps = split_steps(text)
    if not steps:
        raise ValueError("Không tìm thấy bước nào trong ai_output.")

    doc = Document()
    for section in doc.sections:
        section.page_width    = Cm(21)
        section.page_height   = Cm(29.7)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)

    # Tiêu đề tài liệu
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("HƯỚNG DẪN SỬ DỤNG")
    set_run_font(run, 20, bold=True, color=COLOR_TITLE)
    set_paragraph_spacing(title_para, before_pt=0, after_pt=18)

    for step_title, content in steps:
        img_path = extract_image_path(content)
        pixel    = extract_pixel(content)
        clean    = clean_meta_lines(content)

        # Tiêu đề bước
        h = doc.add_paragraph()
        set_run_font(h.add_run(step_title), 14, bold=True, color=COLOR_TITLE)
        set_paragraph_spacing(h, before_pt=14, after_pt=6)
        add_paragraph_border_bottom(h, color_hex=COLOR_BORDER, size=6)

        # Nội dung
        if clean:
            for line in clean.split("\n"):
                s = line.strip()
                if not s:
                    continue
                if re.match(r"^[-*]\s+", s):
                    p = doc.add_paragraph(style="List Bullet")
                    add_inline_runs(p, re.sub(r"^[-*]\s+", "", s))
                    set_paragraph_spacing(p, after_pt=4)
                elif s.startswith(">"):
                    p = doc.add_paragraph()
                    add_inline_runs(p, s[1:].strip(), base_size=10, base_color=COLOR_NOTE)
                    set_paragraph_indent(p, left_cm=1.0)
                    add_paragraph_border_left(p, color_hex=COLOR_BORDER)
                    set_paragraph_spacing(p, after_pt=4)
                else:
                    p = doc.add_paragraph()
                    add_inline_runs(p, s)
                    set_paragraph_spacing(p, after_pt=4)

        # Hình ảnh
        if img_path:
            label = doc.add_paragraph()
            set_run_font(label.add_run("Hình ảnh minh họa:"), 10, bold=True, color=COLOR_NOTE)
            set_paragraph_spacing(label, before_pt=8, after_pt=4)
            add_scaled_image(doc, resolve_image_file(img_path, img_base))

        # Pixel
        if pixel:
            p = doc.add_paragraph()
            set_run_font(p.add_run("Vị trí tương tác: "), 9, bold=True, color=COLOR_NOTE)
            set_run_font(p.add_run(pixel), 9, mono=True, color=COLOR_NOTE)
            set_paragraph_indent(p, left_cm=0.5)
            set_paragraph_spacing(p, after_pt=4)

        # Divider
        hr = doc.add_paragraph()
        set_paragraph_spacing(hr, before_pt=10, after_pt=10)
        add_paragraph_border_bottom(hr, color_hex="DDDDDD", size=4)

    doc.save(str(out))
    print(f"✅ Đã xuất Word: {out}")
    return out


# ── Standalone ───────────────────────────────────────────────────────────────
if __name__ == "__main__":
    build_docx()